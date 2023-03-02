import requests
import docx
from bs4 import BeautifulSoup

# Set the URL to scrape
url = 'https://example.com'

# Send a request to get the HTML content
response = requests.get(url)
html = response.text

# Parse the HTML using BeautifulSoup
soup = BeautifulSoup(html, 'html.parser')

# Find the information to extract
info = soup.find('div', {'class': 'info'})

# Write the information to a Word document
doc = docx.Document()
doc.add_heading('Information', 0)
if info is not None:
    doc.add_paragraph(info.text)
else:
    doc.add_paragraph('No information found')
doc.save('information.docx')